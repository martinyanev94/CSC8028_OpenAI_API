import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import openai
import config
import requests
from PIL import Image, ImageTk
import io
from docx import Document

# Set your OpenAI API key
openai.api_key = config.API_KEY

class AIApp(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("AI Application Suite")
        self.geometry("800x600")

        container = tk.Frame(self)
        container.pack(side="top", fill="both", expand=True)
        container.grid_rowconfigure(0, weight=1)
        container.grid_columnconfigure(0, weight=1)

        self.frames = {}

        for F in (HomePage, EssayGenerationPage, ImageGenerationPage, ChatBotPage, AudioTranscriptionPage, TextSummarizationPage):
            frame = F(container, self)
            self.frames[F] = frame
            frame.grid(row=0, column=0, sticky="nsew")

        self.show_frame(HomePage)

    def show_frame(self, page):
        """Show a frame for the given page."""
        frame = self.frames[page]
        frame.tkraise()


class HomePage(tk.Frame):
    def __init__(self, parent, controller):
        super().__init__(parent)
        tk.Label(self, text="Welcome to the AI Application Suite!", font=("Helvetica", 18)).pack(pady=20)
        tk.Label(self, text="Select a tool below to get started.", font=("Helvetica", 14)).pack(pady=10)

        buttons = [
            ("Essay Generation", EssayGenerationPage),
            ("Image Generation", ImageGenerationPage),
            ("ChatBot", ChatBotPage),
            ("Audio Transcription", AudioTranscriptionPage),
            ("Text Summarization", TextSummarizationPage),
        ]

        for text, page in buttons:
            ttk.Button(self, text=text, command=lambda p=page: controller.show_frame(p)).pack(pady=10)


class EssayGenerationPage(tk.Frame):
    def __init__(self, parent, controller):
        super().__init__(parent)
        tk.Label(self, text="Essay Generation", font=("Helvetica", 18)).pack(pady=20)
        tk.Label(self, text="Enter a topic to generate an essay:").pack()

        self.topic_entry = tk.Entry(self, width=50)
        self.topic_entry.pack(pady=10)

        self.result_display = tk.Text(self, height=10, width=60)
        self.result_display.pack(pady=10)

        ttk.Button(self, text="Generate Essay", command=self.generate_essay).pack(pady=5)
        ttk.Button(self, text="Export to Word", command=self.export_to_word).pack(pady=5)
        ttk.Button(self, text="Back to Home", command=lambda: controller.show_frame(HomePage)).pack(pady=20)

    def generate_essay(self):
        topic = self.topic_entry.get()
        essay = self.generate_essay_from_openai(topic)
        self.result_display.delete("1.0", "end")
        self.result_display.insert("1.0", essay)

    def generate_essay_from_openai(self, topic):
        try:
            messages = [
                {"role": "system", "content": "You are an AI that can generate detailed essays."},
                {"role": "user", "content": f"Write a detailed essay on the following topic: {topic}"}
            ]

            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=messages,
                max_tokens=1024,
                temperature=0.7
            )

            return response['choices'][0]['message']['content'].strip()
        except Exception as e:
            return f"Error generating essay: {str(e)}"

    def export_to_word(self):
        essay = self.result_display.get("1.0", "end-1c")
        if not essay.strip():
            messagebox.showwarning("No Content", "No essay content to export.")
            return

        doc = Document()
        doc.add_heading("Generated Essay", 0)
        doc.add_paragraph(essay)
        file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Documents", "*.docx")])
        if file_path:
            doc.save(file_path)
            messagebox.showinfo("Success", "Essay exported to Word successfully!")


class ImageGenerationPage(tk.Frame):
    def __init__(self, parent, controller):
        super().__init__(parent)
        tk.Label(self, text="Image Generation", font=("Helvetica", 18)).pack(pady=20)
        tk.Label(self, text="Describe the image you want to generate:").pack()

        self.description_entry = tk.Entry(self, width=50)
        self.description_entry.pack(pady=10)

        self.result_label = tk.Label(self, text="Generated image will appear here.")
        self.result_label.pack(pady=10)

        ttk.Button(self, text="Generate Image", command=self.generate_image).pack(pady=5)
        ttk.Button(self, text="Download Image", command=self.download_image).pack(pady=5)
        ttk.Button(self, text="Back to Home", command=lambda: controller.show_frame(HomePage)).pack(pady=20)

        self.image_data = None  # Variable to store the generated image data

    def generate_image(self):
        description = self.description_entry.get()
        if not description.strip():
            messagebox.showwarning("Input Required", "Please enter a description for the image.")
            return

        try:
            response = openai.Image.create(
                prompt=description,
                n=1,
                size="1024x1024"
            )

            image_url = response['data'][0]['url']
            self.display_image(image_url)

        except Exception as e:
            messagebox.showerror("Error", f"Error generating image: {str(e)}")

    def display_image(self, image_url):
        try:
            response = requests.get(image_url)
            self.image_data = response.content
            image = Image.open(io.BytesIO(self.image_data))
            image = image.resize((300, 300))

            img_tk = ImageTk.PhotoImage(image)

            self.result_label.config(image=img_tk)
            self.result_label.image = img_tk

        except Exception as e:
            messagebox.showerror("Error", f"Error displaying image: {str(e)}")

    def download_image(self):
        if not self.image_data:
            messagebox.showwarning("No Image", "No image available to download.")
            return

        file_path = filedialog.asksaveasfilename(defaultextension=".png", filetypes=[("PNG Image", "*.png")])
        if file_path:
            try:
                with open(file_path, 'wb') as f:
                    f.write(self.image_data)
                messagebox.showinfo("Success", "Image downloaded successfully!")
            except Exception as e:
                messagebox.showerror("Error", f"Error saving image: {str(e)}")


class ChatBotPage(tk.Frame):
    def __init__(self, parent, controller):
        super().__init__(parent)
        tk.Label(self, text="ChatBot", font=("Helvetica", 18)).pack(pady=20)

        self.conversation_display = tk.Text(self, height=15, width=60)
        self.conversation_display.pack(pady=10)

        self.message_entry = tk.Entry(self, width=50)
        self.message_entry.pack(pady=10)

        ttk.Button(self, text="Send", command=self.send_message).pack(pady=5)
        ttk.Button(self, text="Back to Home", command=lambda: controller.show_frame(HomePage)).pack(pady=20)

    def send_message(self):
        user_message = self.message_entry.get()
        self.conversation_display.insert("end", f"User: {user_message}\n")
        self.message_entry.delete(0, "end")

        self.get_bot_response(user_message)

    def get_bot_response(self, user_message):
        try:
            messages = [
                {"role": "system", "content": "You are a helpful assistant."},
                {"role": "user", "content": user_message}
            ]

            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=messages,
                max_tokens=150,
                temperature=0.7
            )

            bot_message = response['choices'][0]['message']['content'].strip()
            self.conversation_display.insert("end", f"Bot: {bot_message}\n")
            self.conversation_display.yview("end")

        except Exception as e:
            self.conversation_display.insert("end", f"Error: {str(e)}\n")


class AudioTranscriptionPage(tk.Frame):
    def __init__(self, parent, controller):
        super().__init__(parent)
        tk.Label(self, text="Audio Transcription", font=("Helvetica", 18)).pack(pady=20)

        ttk.Button(self, text="Upload Audio File", command=self.upload_audio).pack(pady=10)

        self.transcription_display = tk.Text(self, height=10, width=60)
        self.transcription_display.pack(pady=10)

        ttk.Button(self, text="Back to Home", command=lambda: controller.show_frame(HomePage)).pack(pady=20)

    def upload_audio(self):
        file_path = filedialog.askopenfilename(filetypes=[("Audio Files", "*.mp3")])
        if not file_path:
            return

        try:
            with open(file_path, 'rb') as audio_file:
                transcription = openai.Audio.transcribe(
                    model="whisper-1",
                    file=audio_file
                )

            self.transcription_display.delete("1.0", "end")
            self.transcription_display.insert("1.0", transcription['text'])

        except Exception as e:
            self.transcription_display.delete("1.0", "end")
            self.transcription_display.insert("1.0", f"Error: {str(e)}")


class TextSummarizationPage(tk.Frame):
    def __init__(self, parent, controller):
        super().__init__(parent)
        tk.Label(self, text="Text Summarization", font=("Helvetica", 18)).pack(pady=20)

        tk.Label(self, text="Enter the text to summarize:").pack()

        self.text_entry = tk.Text(self, height=10, width=60)
        self.text_entry.pack(pady=10)

        self.summarized_text_display = tk.Text(self, height=10, width=60)
        self.summarized_text_display.pack(pady=10)

        ttk.Button(self, text="Summarize Text", command=self.summarize_text).pack(pady=5)
        ttk.Button(self, text="Back to Home", command=lambda: controller.show_frame(HomePage)).pack(pady=20)

    def summarize_text(self):
        text_to_summarize = self.text_entry.get("1.0", "end-1c")
        if not text_to_summarize.strip():
            messagebox.showwarning("Input Required", "Please enter text to summarize.")
            return

        try:
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "You are an AI that can summarize text."},
                    {"role": "user", "content": f"Please summarize the following text: {text_to_summarize}"}
                ],
                max_tokens=200,
                temperature=0.7
            )

            summarized_text = response['choices'][0]['message']['content'].strip()
            self.summarized_text_display.delete("1.0", "end")
            self.summarized_text_display.insert("1.0", summarized_text)

        except Exception as e:
            self.summarized_text_display.delete("1.0", "end")
            self.summarized_text_display.insert("1.0", f"Error: {str(e)}")


if __name__ == "__main__":
    app = AIApp()
    app.mainloop()
